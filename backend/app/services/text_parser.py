import pdfplumber
from docx import Document
import re
from typing import Optional, Dict, Any
import logging

logger = logging.getLogger(__name__)

class TextParser:
    """Service for parsing text from various file formats"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF files using pdfplumber"""
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX files using python-docx"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                    
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
            
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\.\,\-\(\)\+\#\@]', ' ', text)
        
        # Normalize spacing
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    @staticmethod
    def extract_contact_info(text: str) -> Dict[str, Optional[str]]:
        """Extract contact information from resume text"""
        contact_info = {
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None
        }
        
        # Email extraction
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info["email"] = email_match.group()
        
        # Phone extraction
        phone_pattern = r'(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact_info["phone"] = phone_match.group()
        
        # LinkedIn extraction
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact_info["linkedin"] = linkedin_match.group()
        
        # GitHub extraction
        github_pattern = r'github\.com/[\w-]+'
        github_match = re.search(github_pattern, text, re.IGNORECASE)
        if github_match:
            contact_info["github"] = github_match.group()
        
        return contact_info
    
    @staticmethod
    def extract_sections(text: str) -> Dict[str, str]:
        """Extract different sections from resume text"""
        sections = {
            "education": "",
            "experience": "",
            "skills": "",
            "certifications": "",
            "projects": ""
        }
        
        # Common section headers patterns
        patterns = {
            "education": r'(education|academic|qualification|degree)',
            "experience": r'(experience|employment|work|career|professional)',
            "skills": r'(skills|technical|competenc|technolog)',
            "certifications": r'(certification|certificate|license)',
            "projects": r'(projects|portfolio|work samples)'
        }
        
        text_lower = text.lower()
        lines = text.split('\n')
        
        current_section = None
        section_content = []
        
        for line in lines:
            line_lower = line.strip().lower()
            if not line_lower:
                continue
                
            # Check if line is a section header
            section_found = False
            for section, pattern in patterns.items():
                if re.search(pattern, line_lower) and len(line_lower) < 50:
                    if current_section and section_content:
                        sections[current_section] = '\n'.join(section_content)
                    current_section = section
                    section_content = []
                    section_found = True
                    break
            
            if not section_found and current_section:
                section_content.append(line.strip())
        
        # Add the last section
        if current_section and section_content:
            sections[current_section] = '\n'.join(section_content)
        
        return sections